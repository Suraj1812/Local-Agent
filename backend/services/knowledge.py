import json
from io import BytesIO
from typing import Any, Dict, List

from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile
from pypdf import PdfReader
from sqlalchemy.orm import Session

from database.models import Document, DocumentChunk
from services.embedding import cosine_similarity, embed_text
from services.text_utils import chunk_text

MAX_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_TEXT_CHARS = 750_000
MAX_SEARCH_ROWS = 1500


async def extract_text(file: UploadFile) -> str:
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File is too large")

    name = (file.filename or "").lower()
    mime = file.content_type or ""

    try:
        if "pdf" in mime or name.endswith(".pdf"):
            reader = PdfReader(BytesIO(content))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif "wordprocessingml" in mime or name.endswith(".docx"):
            document = DocxDocument(BytesIO(content))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        else:
            text = content.decode("utf-8", errors="ignore")
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Could not read uploaded file") from exc

    text = text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="No readable text found")
    return text[:MAX_TEXT_CHARS]


async def add_document(db: Session, file: UploadFile) -> Dict[str, Any]:
    text = await extract_text(file)
    document = Document(
        filename=file.filename or "upload",
        mime_type=file.content_type or "application/octet-stream",
        content=text,
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    chunks = chunk_text(text)
    for index, chunk in enumerate(chunks):
        db.add(
            DocumentChunk(
                document_id=document.id,
                chunk_index=index,
                content=chunk,
                embedding_json=json.dumps(embed_text(chunk)),
            )
        )
    db.commit()
    return {"id": document.id, "filename": document.filename, "mime_type": document.mime_type, "chunks": len(chunks)}


def list_documents(db: Session) -> List[Dict[str, Any]]:
    rows = db.query(Document).order_by(Document.created_at.desc()).all()
    output = []
    for row in rows:
        chunks = db.query(DocumentChunk).filter(DocumentChunk.document_id == row.id).count()
        output.append(
            {
                "id": row.id,
                "filename": row.filename,
                "mime_type": row.mime_type,
                "chunks": chunks,
                "created_at": row.created_at,
            }
        )
    return output


def search_knowledge(db: Session, query: str, limit: int = 6) -> List[Dict[str, Any]]:
    query = " ".join(query.strip().split())
    if len(query) < 2:
        return []
    query_vector = embed_text(query)
    rows = (
        db.query(DocumentChunk, Document)
        .join(Document, Document.id == DocumentChunk.document_id)
        .order_by(DocumentChunk.id.desc())
        .limit(MAX_SEARCH_ROWS)
        .all()
    )
    ranked = []
    for chunk, document in rows:
        score = cosine_similarity(query_vector, json.loads(chunk.embedding_json))
        if score > 0:
            ranked.append(
                {
                    "document_id": document.id,
                    "filename": document.filename,
                    "chunk_index": chunk.chunk_index,
                    "content": chunk.content,
                    "score": score,
                }
            )
    return sorted(ranked, key=lambda item: item["score"], reverse=True)[:limit]
